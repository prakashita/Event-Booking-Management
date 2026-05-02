"""
IQAC SSR / NAAC Word Document Generator
=========================================
Generates an A4 .docx following the same section order as ssr_pdf.py:

  Cover Page
  1. Executive Summary
  2. Profile of the University
  3. Extended Profile
  4. Quality Indicator Framework

Dependencies: python-docx >= 1.1
"""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _d(v) -> dict:
    return v if isinstance(v, dict) else {}


def _s(data, *keys, default="") -> str:
    cur = data
    for k in keys:
        if not isinstance(cur, dict):
            return str(default)
        cur = cur.get(k, default)
    return str(cur or default).strip()


def _rows(data, *keys) -> list:
    cur = data
    for k in keys:
        if not isinstance(cur, dict):
            return []
        cur = cur.get(k, [])
    return cur if isinstance(cur, list) else []


def _yes_no_marks(value) -> tuple[str, str]:
    normalized = str(value or "").strip().lower()
    if normalized in {"yes", "y", "true", "1"}:
        return "X", ""
    if normalized in {"no", "n", "false", "0"}:
        return "", "X"
    return "", ""


# Typography
FONT_MAIN = "Times New Roman"


def _run_fmt(run, bold=False, italic=False, size=10, color=None):
    run.font.name    = FONT_MAIN
    run.font.bold    = bold
    run.font.italic  = italic
    run.font.size    = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _heading(doc: Document, text: str, level: int = 1, keep_with_next=False):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    _run_fmt(run, bold=True, size=14 - level)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = keep_with_next
    return p


def _para(doc: Document, text: str, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT, size=10, space_after=4,
          keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = keep_with_next
    run = p.add_run(text)
    _run_fmt(run, bold=bold, italic=italic, size=size)
    return p


def _repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def _cant_split_row(row):
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)


def _hr(doc: Document):
    """Add a simple horizontal paragraph border as a divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_table(doc: Document, rows_data: list[list[str]],
               col_widths_cm: list[float],
               header_row: bool = True,
               repeat_header: bool = True):
    """Add a bordered table. First row is bolded if header_row=True."""
    if not rows_data:
        return None
    n_cols = len(rows_data[0])
    tbl = doc.add_table(rows=len(rows_data), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.autofit = False
    for r_idx, row in enumerate(rows_data):
        tr = tbl.rows[r_idx]
        if header_row and repeat_header and r_idx == 0:
            _repeat_table_header(tr)
        if r_idx <= 1:
            _cant_split_row(tr)
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            run  = para.add_run(str(cell_text or ""))
            run.font.name  = FONT_MAIN
            run.font.size  = Pt(9)
            run.font.bold  = (header_row and r_idx == 0)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            # Column width
            if c_idx < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[c_idx])
    doc.add_paragraph()  # spacer after table
    return tbl


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------
def _build_cover(doc: Document, sections_data: dict, meta: dict):
    p_data = _d(sections_data.get("university_profile"))
    basic  = _d(p_data.get("basic_information"))
    inst   = _s(basic, "name") or "Institution"

    generated_by = meta.get("generated_by") or "IQAC Administrator"
    generated_at = meta.get("generated_at")
    if isinstance(generated_at, datetime):
        ts = generated_at.strftime("%d %B %Y, %I:%M %p") + " (UTC)"
    else:
        ts = datetime.utcnow().strftime("%d %B %Y, %I:%M %p") + " (UTC)"

    for _ in range(3):
        doc.add_paragraph()

    _para(doc, "SELF STUDY REPORT (SSR)", bold=True, size=18,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _para(doc, "Submitted to", italic=True, size=13,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "National Assessment and Accreditation Council (NAAC)",
          bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "Bengaluru - 560 072, Karnataka",
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    _hr(doc)
    _para(doc, inst, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    addr_parts = [_s(basic, k) for k in ("address", "city", "state", "pin") if _s(basic, k)]
    detail_rows = [
        ["Address",      ", ".join(addr_parts) or "-"],
        ["Website",      _s(basic, "website") or "-"],
        ["Generated by", generated_by],
        ["Generated on", ts],
        ["System",       "IQAC SSR / NAAC Data Entry Portal"],
    ]
    _add_table(doc, detail_rows, [4.5, 12.0])
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 1 - Executive Summary
# ---------------------------------------------------------------------------
def _build_executive_summary(doc: Document, data: dict):
    es = _d(data.get("executive_summary"))
    _heading(doc, "1. Executive Summary", level=1)
    _para(doc, (
        "Every Higher Education Institution (HEI) applying for the Accreditation and "
        "Assessment (A&A) process shall prepare an Executive Summary highlighting the main "
        "features of the Institution. The Executive Summary shall not be more than 5000 words."
    ), italic=True, size=9, space_after=8)

    fields = [
        ("1.1 Introductory Note on the Institution", "introductory_note"),
        ("1.2 Criterion-wise Summary",               "criteria_summary"),
        ("1.3 SWOC Analysis",                        "swoc_analysis"),
        ("1.4 Additional Information",               "additional_information"),
        ("1.5 Overall Conclusive Explication",       "conclusive_explication"),
    ]
    for heading, key in fields:
        _heading(doc, heading, level=2)
        content = _s(es, key)
        if content:
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    _para(doc, line, space_after=4)
        else:
            _para(doc, "(Not filled)", italic=True, size=9)


# ---------------------------------------------------------------------------
# Section 2 - Profile of the University
# ---------------------------------------------------------------------------
def _build_profile(doc: Document, data: dict):
    p = _d(data.get("university_profile"))
    doc.add_page_break()
    _heading(doc, "2.  Profile of the University", level=1)

    basic = _d(p.get("basic_information"))
    _heading(doc, "Basic Information", level=2)
    _add_table(doc, [
        ["Name and Address of the University", "", "", ""],
        ["Name",    _s(basic, "name"),    "",           ""],
        ["Address", _s(basic, "address"), "",           ""],
        ["City",    _s(basic, "city"),    "Pin",        _s(basic, "pin")],
        ["State",   _s(basic, "state"),   "Website",    _s(basic, "website")],
    ], [2.5, 6.5, 2.5, 6.5], header_row=False)

    _heading(doc, "Contacts for Communication", level=2)
    contacts_raw = p.get("contacts")
    ct_rows = [["Designation", "Name", "Telephone (STD)", "Mobile", "Fax", "Email"]]
    if isinstance(contacts_raw, list):
        for c in contacts_raw:
            if isinstance(c, dict):
                ct_rows.append([_s(c, "designation"), _s(c, "name"),
                                  _s(c, "telephone"), _s(c, "mobile"),
                                  _s(c, "fax"),       _s(c, "email")])
    elif isinstance(contacts_raw, dict):
        for rk, lbl in [("head_of_institution", "Head of Institution"),
                         ("iqac_coordinator",    "IQAC Coordinator")]:
            c = _d(contacts_raw.get(rk))
            ct_rows.append([_s(c, "designation") or lbl, _s(c, "name"),
                              _s(c, "telephone"), _s(c, "mobile") or _s(c, "phone"),
                              _s(c, "fax"),       _s(c, "email")])
    if len(ct_rows) == 1:
        ct_rows.append(["(No contacts entered)", "", "", "", "", ""])
    _add_table(doc, ct_rows, [3.2, 3.0, 3.0, 2.5, 1.8, 4.3])

    inst  = _d(p.get("institution"))
    estab = _d(p.get("establishment"))
    recog = _d(p.get("recognition"))

    _heading(doc, "Nature / Type / Establishment", level=2)
    _add_table(doc, [
        ["Nature of University",   _s(inst, "nature"),  "Institution Status",  _s(inst, "status")],
        ["Type of University",     _s(inst, "type"),    "Type of University",  _s(inst, "type")],
        ["Establishment Date",     _s(estab, "establishment_date"), "", ""],
        ["Status Prior (if appl)", _s(estab, "status_prior"), "", ""],
    ], [2.8, 6.2, 2.8, 6.2], header_row=False)

    _heading(doc, "Recognition Details", level=2)
    ugc_2f  = _s(recog, "ugc_2f_date")  or _s(recog, "ugc_recognition_date")
    ugc_12b = _s(recog, "ugc_12b_date") or _s(recog, "section_12b")
    _add_table(doc, [
        ["Date of Recognition as University by UGC or Any Other National Agency", ""],
        ["Under Section",    "Date"],
        ["2f of UGC",        ugc_2f],
        ["12B of UGC",       ugc_12b],
    ], [7.0, 11.0])

    _heading(doc, "University with Potential for Excellence (UPE)", level=2, keep_with_next=True)
    upe_yes, upe_no = _yes_no_marks(_s(p, "upe_recognized"))
    _add_table(doc, [
        ["Is the University Recognised as UPE by UGC?", "Yes", "No"],
        ["", upe_yes, upe_no],
    ], [13.0, 2.5, 2.5])

    _heading(doc, "Location, Area and Activity of Campus", level=2)
    camp_rows = [["Campus Type", "Address", "Location", "Area (Acres)",
                   "Built-up (sq.m.)", "Programmes", "Estd. Date", "Recog. Date"]]
    for camp in (_rows(p, "campuses") or [{}]):
        camp_rows.append([
            _s(camp, "campus_type"),       _s(camp, "address"),
            _s(camp, "location"),          _s(camp, "campus_area_acres"),
            _s(camp, "built_up_area_sq_mts"), _s(camp, "programmes_offered"),
            _s(camp, "establishment_date"),_s(camp, "recognition_date"),
        ])
    _add_table(doc, camp_rows, [2.0, 2.5, 2.5, 2.0, 2.0, 2.5, 2.5, 2.5])

    _heading(doc, "Academic Information", level=2)
    ai = _d(p.get("academic_information"))

    aff_rows = _rows(ai, "affiliated_institutions")
    aff1 = [["College Type", "Permanent Affiliation", "Temporary Affiliation"]]
    for row in (aff_rows or [{}]):
        aff1.append([_s(row, "college_type"),
                      _s(row, "permanent_affiliation"), _s(row, "temporary_affiliation")])
    _add_table(doc, aff1, [6.0, 5.0, 5.0])

    ct_rows = _rows(ai, "college_type_affiliations") or [
        {"college_type": "Education/Teachers Training"},
        {"college_type": "Business administration/Commerce/Management"},
        {"college_type": "Universal/Common to all Disciplines"},
    ]
    aff2 = [["Type of Colleges", "Permanent", "Temporary", "Total"]]
    for row in ct_rows:
        aff2.append([_s(row, "college_type"),
                      _s(row, "permanent") or _s(row, "permanent_affiliation"),
                      _s(row, "temporary") or _s(row, "temporary_affiliation"),
                      _s(row, "total")])
    _add_table(doc, aff2, [8.5, 3.0, 3.0, 3.5])

    _para(doc, "Furnish the Details of Colleges under University", bold=True, size=9)
    college_details = _rows(ai, "college_details") or [
        {"label": "Constituent Colleges"},
        {"label": "Affiliated Colleges"},
        {"label": "Colleges Under 2(f)"},
        {"label": "Colleges Under 2(f) and 12B"},
        {"label": "NAAC Accredited Colleges"},
        {"label": "Colleges with Potential for Excellence (UGC)"},
        {"label": "Autonomous Colleges"},
        {"label": "Colleges with Postgraduate Departments"},
        {"label": "Colleges with Research Departments"},
        {"label": "University Recognized Research Institutes/Centers"},
    ]
    cd_data = [[_s(r, "label"), _s(r, "value")] for r in college_details]
    _add_table(doc, cd_data, [13.0, 5.0], header_row=False)

    sra_yes, sra_no = _yes_no_marks(_s(ai, "sra_recognized"))
    _add_table(doc, [
        ["Is the University Offering Programmes Recognised by any SRA?", "Yes", "No"],
        ["", sra_yes, sra_no],
    ], [13.0, 2.5, 2.5])

    # Teaching Staff
    _heading(doc, "Details of Teaching & Non-Teaching Staff", level=2)
    _para(doc, "Teaching Faculty", bold=True, size=10)
    staff   = _d(p.get("staff"))
    tf_rows = _rows(staff, "teaching")
    ROLES_T  = ["Professor", "Associate Professor", "Assistant Professor"]
    GENDERS  = ["Male", "Female", "Others", "Total"]
    STATUSES = ["Sanctioned", "Recruited", "Yet to Recruit", "On Contract"]
    tf_hdr = [""] + [f"{r}\n{g}" for r in ROLES_T for g in GENDERS] + ["Total"]
    tf_table = [tf_hdr]
    for row in (tf_rows if tf_rows else [{"status": s} for s in STATUSES]):
        cells = [_s(row, "status")]
        for role in ROLES_T:
            for g in GENDERS:
                cells.append(_s(row, f"{role}_{g}"))
        cells.append(_s(row, "Total"))
        tf_table.append(cells)
    _add_table(doc, tf_table, [2.2] + [1.1]*12 + [0.8])

    for grp_lbl, grp_key in [("Non-Teaching Staff", "non_teaching"),
                               ("Technical Staff",   "technical")]:
        _para(doc, grp_lbl, bold=True, size=10)
        g_rows = _rows(staff, grp_key)
        nt = [["", "Male", "Female", "Others", "Total"]]
        for row in (g_rows if g_rows else [{"status": s} for s in STATUSES]):
            nt.append([_s(row, "status"),
                        _s(row, "Male"), _s(row, "Female"),
                        _s(row, "Others"), _s(row, "Total")])
        _add_table(doc, nt, [5.0, 3.0, 3.0, 3.0, 3.0])

    # Qualification Details
    _heading(doc, "Qualification Details of the Teaching Staff", level=2)
    qual = _d(p.get("qualification_details"))
    QUALS = ["D.sc/D.Litt", "Ph.D.", "M.Phil.", "PG"]
    qualification_widths = [2.6] + [1.25] * 9 + [2.1]
    for grp_lbl, grp_key in [("Permanent Teachers", "permanent_teachers"),
                               ("Temporary Teachers", "temporary_teachers"),
                               ("Part Time Teachers",  "part_time_teachers")]:
        _para(doc, grp_lbl, bold=True, size=10, keep_with_next=True)
        q_rows = _rows(qual, grp_key)
        qh = ["Highest Qualification"] + [f"{r}\n{g}"
              for r in ["Professor", "Associate Professor", "Assistant Professor"]
              for g in ("Male", "Female", "Others")] + ["Total"]
        q_data = [qh]
        for row in (q_rows if q_rows else [{"qualification": q} for q in QUALS]):
            cells = [_s(row, "qualification")]
            for role_k in ["Professor", "Associate Professor", "Assistant Professor"]:
                for g in ("Male", "Female", "Others"):
                    cells.append(_s(row, f"{role_k}_{g}"))
            cells.append(_s(row, "Total"))
            q_data.append(cells)
        _add_table(doc, q_data, qualification_widths)

    # Distinguished Academicians
    _heading(doc, "Distinguished Academicians Appointed", level=2, keep_with_next=True)
    da_rows = _rows(p, "distinguished_academicians") or [
        {"role": r} for r in ("Emeritus Professor", "Adjunct Professor", "Visiting Professor")]
    da_data = [["", "Male", "Female", "Others", "Total"]]
    for row in da_rows:
        da_data.append([_s(row, "role"), _s(row, "male"),
                         _s(row, "female"), _s(row, "others"), _s(row, "total")])
    _add_table(doc, da_data, [5.0, 3.0, 3.0, 3.0, 3.0])

    # Chairs
    _heading(doc, "Chairs Instituted by the University", level=2, keep_with_next=True)
    chairs = _rows(p, "chairs")
    ch_data = [["Sl.No", "Department", "Chair Name", "Sponsor"]]
    for i, row in enumerate(chairs if chairs else [{}], 1):
        ch_data.append([str(i), _s(row, "department"), _s(row, "chair"), _s(row, "sponsor")])
    _add_table(doc, ch_data, [1.5, 5.0, 5.0, 6.5])

    # Student Enrolment
    _heading(doc, "Student Enrolment", level=2, keep_with_next=True)
    enrol = _rows(p, "student_enrolment")
    se_data = [["Programme", "Gender", "From State", "From Other States",
                 "NRI", "Foreign", "Total"]]
    if enrol:
        prev_prog = None
        for row in enrol:
            prog = _s(row, "programme")
            se_data.append([prog if prog != prev_prog else "",
                              _s(row, "gender"),
                              _s(row, "from_state"),
                              _s(row, "from_other_states"),
                              _s(row, "nri"), _s(row, "foreign"), _s(row, "total")])
            prev_prog = prog
    else:
        for prog in ("PG", "UG", "PG Diploma"):
            for gender in ("Male", "Female", "Others"):
                se_data.append([prog if gender == "Male" else "",
                                  gender, "", "", "", "", ""])
    _add_table(doc, se_data, [2.5, 1.5, 3.5, 3.5, 2.0, 2.0, 2.0])

    # Integrated Programmes
    _heading(doc, "Integrated Programmes", level=2, keep_with_next=True)
    integ = _d(p.get("integrated_programmes"))
    ip_yes, ip_no = _yes_no_marks(_s(integ, "offered"))
    _add_table(doc, [
        ["Does the university offer any integrated programmes?", "Yes", "No"],
        ["", ip_yes, ip_no],
        ["Total number of integrated programme", _s(integ, "total_programmes")],
    ], [11.0, 3.0, 3.0], header_row=False)
    ip_enrol = _rows(integ, "enrolment")
    ip_data  = [["Gender", "From State", "From Other States", "NRI", "Foreign", "Total"]]
    for row in (ip_enrol if ip_enrol else [{"gender": g} for g in ("Male","Female","Others")]):
        ip_data.append([_s(row, "gender"), _s(row, "from_state"),
                         _s(row, "from_other_states"), _s(row, "nri"),
                         _s(row, "foreign"), _s(row, "total")])
    _add_table(doc, ip_data, [3.0, 4.0, 4.0, 2.5, 2.5, 2.0])

    # HRDC
    _heading(doc, "UGC Human Resource Development Centre (HRDC)", level=2, keep_with_next=True)
    hrdc = _d(p.get("hrdc"))
    hrdc_data = [
        ["Year of Establishment",                                   _s(hrdc, "year_of_establishment")],
        ["Number of UGC Orientation Programmes",                    _s(hrdc, "orientation_programmes")],
        ["Number of UGC Refresher Course",                          _s(hrdc, "refresher_courses")],
        ["Number of University's own Programmes",                   _s(hrdc, "own_programmes")],
        ["Total Number of Programmes Conducted (last five years)",  _s(hrdc, "total_programmes_last_five_years")],
    ]
    _add_table(doc, hrdc_data, [11.0, 7.0], header_row=False)

    # Department Reports
    _heading(doc, "Evaluative Report of the Departments", level=2, keep_with_next=True)
    dept_reports = _rows(p, "department_reports")
    dr_data = [["Department Name", "Report Reference"]]
    for row in (dept_reports if dept_reports else [{}, {}, {}]):
        dr_data.append([_s(row, "department_name"), _s(row, "report_reference")])
    _add_table(doc, dr_data, [9.0, 9.0])


# ---------------------------------------------------------------------------
# Section 3 - Extended Profile
# ---------------------------------------------------------------------------
YEAR_LABELS_5 = ["2019-20", "2020-21", "2021-22", "2022-23", "2023-24"]


def _metric_year_values_docx(ep: dict, key: str) -> tuple[list[str], list[str]]:
    direct = ep.get(key)
    if isinstance(direct, dict):
        year_labels = direct.get("year_labels") or YEAR_LABELS_5
        values = direct.get("values") or []
    elif isinstance(direct, list):
        year_labels = ep.get("year_labels") or YEAR_LABELS_5
        values = direct
    else:
        metrics_map = _d(ep.get("metrics"))
        raw = metrics_map.get(key)
        year_labels = ep.get("year_labels") or YEAR_LABELS_5
        values = raw if isinstance(raw, list) else []

    if isinstance(values, dict):
        vals = [str(values.get(y, "")) for y in year_labels]
    else:
        vals = [str(v or "") for v in (values or [])]

    while len(vals) < len(year_labels):
        vals.append("")
    vals = vals[:len(year_labels)]
    return [str(y or "") for y in year_labels], vals


def _five_year_table_docx(doc: Document, ep: dict, key: str, row_label: str):
    year_labels, vals = _metric_year_values_docx(ep, key)
    _add_table(doc,
               [["Year"] + list(year_labels), [row_label] + vals],
               [2.5] + [13.5 / len(year_labels)] * len(year_labels))


def _build_extended_profile(doc: Document, data: dict):
    ep = _d(data.get("extended_profile"))
    doc.add_page_break()
    _heading(doc, "3. Extended Profile of the University", level=1)

    metrics_map = _d(ep.get("metrics"))

    def single(key: str) -> str:
        return _s(ep, key) if ep.get(key) else _s(metrics_map, key)

    def metric(label: str, key: str, row_label: str = "Number"):
        _para(doc, label, size=9, space_after=2, keep_with_next=True)
        _five_year_table_docx(doc, ep, key, row_label)

    _para(doc, "1 Programme:", bold=True, size=10, keep_with_next=True)
    metric("1.1 Number of Programmes offered year wise for last five years", "programmes_offered")
    dept_line = "1.2 Number of departments offering academic programmes"
    dept_value = single("departments_offering_programmes")
    if dept_value:
        dept_line = f"{dept_line}: {dept_value}"
    _para(doc, dept_line, size=9, space_after=8)

    _para(doc, "2 Student:", bold=True, size=10, keep_with_next=True)
    metric("2.1 Number of students year wise during the last five years", "students")
    metric("2.2 Number of outgoing / final year students year wise during the last five years", "outgoing_students")
    metric("2.3 Number of students appeared in the University examination year wise during the last five years", "exam_appeared")
    metric("2.4 Number of revaluation applications year wise during the last 5 years", "revaluation_applications")

    _para(doc, "3 Academic:", bold=True, size=10, keep_with_next=True)
    metric("3.1 Number of courses in all Programmes year wise during the last five years", "courses")
    metric("3.2 Number of full time teachers year wise during the last five years", "full_time_teachers")
    metric("3.3 Number of sanctioned posts year wise during the last five years", "sanctioned_posts")

    _para(doc, "4 Institution:", bold=True, size=10, keep_with_next=True)
    metric("4.1 Number of eligible applications received for admissions to all the Programmes year wise during the last five years", "eligible_applications")
    metric("4.2 Number of seats earmarked for reserved category as per GOI/State Govt rule year wise during the last five years", "reserved_seats")
    _para(doc, f"4.3 Total number of classrooms and seminar halls: {single('total_classrooms_seminar_halls') or '________'}", size=9)
    _para(doc, f"4.4 Total number of computers in the campus for academic purpose: {single('total_computers_academic') or '________'}", size=9)
    metric("4.5 Total Expenditure excluding salary year wise during the last five years (INR in Lakhs)", "expenditure_excluding_salary", "Expenditure")


# ---------------------------------------------------------------------------
# Section 4 - Quality Indicator Framework
# ---------------------------------------------------------------------------
def _build_qif(doc: Document, data: dict):
    doc.add_page_break()
    _heading(doc, "4. Quality Indicator Framework (QIF)", level=1)
    _para(doc, "Essential Note:", bold=True, size=10)
    _para(doc, "The SSR has to be filled in an online format available on the NAAC website.", size=10)
    _para(doc, "The QIF given below presents the Metrics under each Key Indicator (KI) for all the seven Criteria.", size=10)
    _para(doc, "While going through the QIF, details are given below each Metric in the form of:", size=10)
    for item in [
        "data required",
        "formula for calculating the information, wherever required, and",
        "File description – for uploading of document where so-ever required.",
    ]:
        _para(doc, f"•  {item}", italic=True, size=9, space_after=1)
    _para(doc, "These will help Institutions in the preparation of their SSR.", size=10, space_after=8)
    _para(doc, (
        "For some Qualitative Metrics (QlM) which seek descriptive data it is specified as to what kind "
        "of information has to be given and how much. It is advisable to keep data accordingly compiled beforehand."
    ), size=10, space_after=8)
    _para(doc, (
        "For the Quantitative Metrics (QnM) wherever formula is given, it must be noted that these are "
        "given merely to inform the HEIs about the manner in which data submitted will be used. That is "
        "the actual online format seeks only data in specified manner which will be processed digitally."
    ), size=10, space_after=8)
    _para(doc, "Metric wise weightage is also given.", size=10, space_after=10)
    _para(doc, (
        "The actual online format may change slightly from the QIF given in this Manual, in order to "
        "bring compatibility with IT design. Observe this carefully while filling up."
    ), size=10)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def generate_ssr_docx(
    sections_data: dict,
    generated_by: str = "",
    generated_at: "datetime | None" = None,
) -> bytes:
    """
    Build and return the SSR Word document as a bytes object.

    Args:
        sections_data: dict with keys executive_summary, university_profile,
                       extended_profile, qif.
        generated_by:  display name / email of the requesting user.
        generated_at:  UTC datetime of generation (defaults to now).
    """
    meta = {
        "generated_by": generated_by or "IQAC Administrator",
        "generated_at": generated_at or datetime.utcnow(),
    }

    doc = Document()

    # Page size: A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = FONT_MAIN
    style.font.size = Pt(10)

    _build_cover(doc, sections_data, meta)
    _build_executive_summary(doc, sections_data)
    _build_profile(doc, sections_data)
    _build_extended_profile(doc, sections_data)
    _build_qif(doc, sections_data)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
